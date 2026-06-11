# SPDX-License-Identifier: Apache-2.0

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from ummaya.tools.documents.models import DocumentFormat, ToolResultStatus
from ummaya.tools.documents.registry import DocumentToolRuntime
from ummaya.tools.documents.tool_defs import (
    DocumentFieldPatch,
    DocumentLocator,
    DocumentPrimitiveRequest,
)


def test_runtime_accepts_issued_bundle_approval_for_multiple_narrative_patches(
    tmp_path: Path,
) -> None:
    source = tmp_path / "self-intro.docx"
    destination = tmp_path / "self-intro-derivative.docx"
    _write_self_intro_docx(source)
    runtime = DocumentToolRuntime(
        session_id="session-authoring-bundle",
        artifact_root=tmp_path / "artifacts",
        enable_default_pdfa_conformance_bridge=False,
    )
    patches = (
        DocumentFieldPatch(
            target_path=f"engine://python-docx/{source.name}/table/1/r1c2",
            value="공공 문서 자동화 도구 개발 지원동기 초안입니다.",
        ),
        DocumentFieldPatch(
            target_path=f"engine://python-docx/{source.name}/table/1/r2c2",
            value="근거 없이는 작성하지 않는 원칙을 세운 성장과정 초안입니다.",
        ),
    )

    unapproved = runtime.document(
        DocumentPrimitiveRequest(
            correlation_id="corr-authoring-bundle-unapproved",
            document=DocumentLocator(path=str(source), expected_format=DocumentFormat.docx),
            operation="fill",
            instruction="지원동기와 성장과정 초안을 반영해줘.",
            patches=patches,
            destination_path=str(destination),
        )
    )

    assert unapproved.status is ToolResultStatus.needs_input
    assert not destination.exists()
    approved_draft_id = _extract_token(
        unapproved.text_summary,
        "approved_draft_id",
    )
    approved_draft_sha256 = _extract_token(
        unapproved.text_summary,
        "approved_draft_sha256",
    )

    approved = runtime.document(
        DocumentPrimitiveRequest(
            correlation_id="corr-authoring-bundle-approved",
            document=DocumentLocator(path=str(source), expected_format=DocumentFormat.docx),
            operation="fill",
            instruction="승인한 자기소개서 초안을 반영해줘.",
            patches=patches,
            destination_path=str(destination),
            approved_draft_id=approved_draft_id,
            approved_draft_sha256=approved_draft_sha256,
        )
    )

    assert approved.status is ToolResultStatus.ok
    assert destination.exists()
    assert approved.diff is not None
    assert len(approved.diff.changes) == 2


def _write_self_intro_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading("자기소개서 양식", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "지원동기"
    table.cell(0, 1).text = ""
    table.cell(1, 0).text = "성장과정"
    table.cell(1, 1).text = ""
    doc.save(path)


def _extract_token(text: str, name: str) -> str:
    pattern = (
        rf"{name}=(draft-[a-f0-9]{{24}})"
        if name == "approved_draft_id"
        else rf"{name}=([a-f0-9]{{64}})"
    )
    match = re.search(pattern, text)
    assert match is not None
    return match.group(1)
